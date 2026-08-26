"""
Only ever called AFTER the approval gate — never wire this directly to a
draft-generation node. See workflow.py for the interrupt() placement.
"""
from docx import Document
from docx.shared import Pt


def export_resume_docx(name: str, tailored_bullets: dict, output_path: str) -> str:
    doc = Document()
    doc.add_heading(name, level=1)

    for section, bullets in tailored_bullets.items():
        doc.add_heading(section, level=2)
        for bullet in bullets:
            p = doc.add_paragraph(bullet, style="List Bullet")
            p.paragraph_format.space_after = Pt(4)

    doc.save(output_path)
    return output_path


def export_cover_letter_docx(cover_letter_text: str, output_path: str) -> str:
    doc = Document()
    for paragraph in cover_letter_text.split("\n\n"):
        doc.add_paragraph(paragraph)
    doc.save(output_path)
    return output_path
